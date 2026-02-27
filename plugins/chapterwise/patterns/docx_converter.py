#!/usr/bin/env python3
"""
docx_converter.py — Convert DOCX files into Chapterwise projects.

JSON-in/JSON-out pattern script.

Input (stdin):
    {"source": "/path/to/novel.docx", "output_dir": "/path/to/output/"}

Optional fields:
    "config": {
        "heading_styles": ["Heading 1", "Heading 2"],
        "chapter_level": 1,
        "preserve_comments": false,
        "preserve_tracked_changes": false,
        "extract_images": true,
        "chapter_pattern": "^Chapter\\s+\\d+",
        "special_sections": ["Prologue", "Epilogue"],
        "skip_sections": [],
        "output_format": "markdown",
        "structure": "flat"
    }

Output (stdout):
    {"event": "complete", "files": 3, "total_words": 12500, "output_dir": "/path/to/output/"}

CLI usage:
    python3 docx_converter.py source.docx output/
"""

import base64
import json
import os
import re
import subprocess
import sys

# ---------------------------------------------------------------------------
# Default configuration
# ---------------------------------------------------------------------------

CONFIG = {
    "heading_styles": ["Heading 1", "Heading 2"],
    "chapter_level": 1,
    "preserve_comments": False,
    "preserve_tracked_changes": False,
    "extract_images": True,
    "chapter_pattern": r"^Chapter\s+\d+",
    "special_sections": ["Prologue", "Epilogue"],
    "skip_sections": [],
    "output_format": "markdown",
    "structure": "flat",
}

# ---------------------------------------------------------------------------
# Script directory — used to locate common/ utilities
# ---------------------------------------------------------------------------

script_dir = os.path.dirname(os.path.abspath(__file__))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def emit(event: dict) -> None:
    """Write a JSON progress event to stdout."""
    print(json.dumps(event, ensure_ascii=False), flush=True)


def word_count(text: str) -> int:
    """Return a rough word count for the given text."""
    return len(text.split())


def run_common(script_name: str, input_data: dict) -> dict:
    """
    Call one of the common/ utility scripts via subprocess.

    Args:
        script_name: filename inside common/, e.g. "chapter_detector.py"
        input_data:  dict to serialise as JSON and pipe to the script's stdin

    Returns:
        Parsed JSON dict from the script's stdout.

    Raises:
        RuntimeError if the subprocess exits non-zero or stdout is not valid JSON.
    """
    script_path = os.path.join(script_dir, "common", script_name)
    result = subprocess.run(
        ["python3", script_path],
        input=json.dumps(input_data),
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip()
        raise RuntimeError(
            f"{script_name} exited with code {result.returncode}: {stderr}"
        )
    stdout = result.stdout.strip()
    if not stdout:
        raise RuntimeError(f"{script_name} produced no output")
    try:
        return json.loads(stdout)
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            f"{script_name} output is not valid JSON: {exc}\nRaw: {stdout[:300]}"
        ) from exc


# ---------------------------------------------------------------------------
# DOCX import — guarded so the module is importable even without python-docx
# ---------------------------------------------------------------------------

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Formatting conversion helpers
# ---------------------------------------------------------------------------

def _run_to_markdown(run) -> str:
    """Convert a single docx Run to its inline markdown representation."""
    text = run.text
    if not text:
        return ""

    # Wrap in order: bold, italic, underline
    if run.bold:
        text = f"**{text}**"
    if run.italic:
        text = f"*{text}*"
    if run.underline:
        text = f"__{text}__"
    return text


def _paragraph_to_markdown(para, cfg: dict) -> str:
    """
    Convert a docx Paragraph to a markdown string.

    Handles:
    - Heading styles  → # / ## / ### prefix
    - List paragraphs → - prefix
    - Inline formatting via runs
    """
    style_name = para.style.name if para.style else ""
    text_parts = [_run_to_markdown(run) for run in para.runs]
    text = "".join(text_parts)

    # Determine heading level from style name
    heading_level = _heading_level_from_style(style_name, para, cfg)

    if heading_level is not None:
        prefix = "#" * heading_level
        return f"{prefix} {text.strip()}"

    # List paragraphs
    if "List" in style_name:
        return f"- {text}"

    return text


def _heading_level_from_style(style_name: str, para, cfg: dict) -> int | None:
    """
    Determine the markdown heading level for a paragraph.

    Priority:
    1. Explicit style match against cfg["heading_styles"]
    2. Built-in "Heading N" style name
    3. Font-size / bold fallback heuristic
    """
    heading_styles = cfg.get("heading_styles", CONFIG["heading_styles"])

    # 1. Check configured heading styles
    for idx, hs in enumerate(heading_styles):
        if style_name == hs:
            return idx + 1

    # 2. Built-in "Heading N" pattern
    m = re.match(r"^Heading\s+(\d+)$", style_name)
    if m:
        return int(m.group(1))

    # 3. Font-size / bold fallback — treat large bold runs as headings
    try:
        if para.runs:
            run = para.runs[0]
            font_size = run.font.size
            is_bold = run.bold
            if font_size is not None and is_bold:
                pt = font_size / Pt(1) if _DOCX_AVAILABLE else font_size / 12700
                if pt >= 18:
                    return 1
                if pt >= 14:
                    return 2
    except Exception:
        pass

    return None


# ---------------------------------------------------------------------------
# Table conversion
# ---------------------------------------------------------------------------

def _table_to_markdown(table) -> str:
    """Convert a docx Table to a GitHub-flavoured markdown table."""
    rows = table.rows
    if not rows:
        return ""

    md_rows = []
    for i, row in enumerate(rows):
        cells = [cell.text.replace("\n", " ").replace("|", "\\|") for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Header separator
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(md_rows)


# ---------------------------------------------------------------------------
# Footnote / endnote extraction
# ---------------------------------------------------------------------------

def _extract_footnotes(doc) -> dict[str, str]:
    """
    Extract footnotes from the document's XML part.

    Returns a dict mapping footnote id (str) → footnote text.
    """
    footnotes: dict[str, str] = {}
    try:
        footnotes_part = doc.part.footnotes_part
        if footnotes_part is None:
            return footnotes
        for fn in footnotes_part._element.findall(qn("w:footnote")):
            fn_id = fn.get(qn("w:id"), "")
            # Skip the separator and continuation footnotes (id -1, 0)
            try:
                if int(fn_id) < 1:
                    continue
            except (ValueError, TypeError):
                continue
            texts = []
            for para in fn.findall(".//" + qn("w:p")):
                for run in para.findall(".//" + qn("w:r")):
                    for t in run.findall(qn("w:t")):
                        texts.append(t.text or "")
            footnotes[fn_id] = "".join(texts).strip()
    except Exception:
        pass
    return footnotes


def _extract_endnotes(doc) -> dict[str, str]:
    """
    Extract endnotes from the document's XML part.

    Returns a dict mapping endnote id (str) → endnote text.
    """
    endnotes: dict[str, str] = {}
    try:
        endnotes_part = doc.part.endnotes_part
        if endnotes_part is None:
            return endnotes
        for en in endnotes_part._element.findall(qn("w:endnote")):
            en_id = en.get(qn("w:id"), "")
            try:
                if int(en_id) < 1:
                    continue
            except (ValueError, TypeError):
                continue
            texts = []
            for para in en.findall(".//" + qn("w:p")):
                for run in para.findall(".//" + qn("w:r")):
                    for t in run.findall(qn("w:t")):
                        texts.append(t.text or "")
            endnotes[en_id] = "".join(texts).strip()
    except Exception:
        pass
    return endnotes


# ---------------------------------------------------------------------------
# Comment extraction
# ---------------------------------------------------------------------------

def _extract_comments(doc) -> dict[str, str]:
    """
    Extract comments from the document's XML part.

    Returns a dict mapping comment id (str) → comment text.
    """
    comments: dict[str, str] = {}
    try:
        comments_part = doc.part.comments_part
        if comments_part is None:
            return comments
        for comment in comments_part._element.findall(qn("w:comment")):
            cid = comment.get(qn("w:id"), "")
            author = comment.get(qn("w:author"), "")
            texts = []
            for para in comment.findall(".//" + qn("w:p")):
                for run in para.findall(".//" + qn("w:r")):
                    for t in run.findall(qn("w:t")):
                        texts.append(t.text or "")
            body = "".join(texts).strip()
            if author:
                comments[cid] = f"{author}: {body}"
            else:
                comments[cid] = body
    except Exception:
        pass
    return comments


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _extract_images(doc, assets_dir: str) -> list[str]:
    """
    Extract embedded images from the document into assets_dir.

    Returns a list of relative image paths (relative to assets_dir's parent).
    """
    os.makedirs(assets_dir, exist_ok=True)
    saved_paths: list[str] = []

    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    ext = os.path.splitext(image_part.partname)[-1] or ".png"
                    filename = f"image_{len(saved_paths) + 1}{ext}"
                    filepath = os.path.join(assets_dir, filename)
                    with open(filepath, "wb") as fh:
                        fh.write(image_part.blob)
                    saved_paths.append(os.path.join("assets", filename))
                except Exception:
                    continue
    except Exception:
        pass

    return saved_paths


# ---------------------------------------------------------------------------
# Document properties
# ---------------------------------------------------------------------------

def _get_doc_properties(doc) -> dict:
    """Extract core document properties (title, author, created date)."""
    props: dict = {
        "title": "",
        "author": "",
        "created": "",
    }
    try:
        core = doc.core_properties
        props["title"] = core.title or ""
        props["author"] = core.author or ""
        created = core.created
        if created is not None:
            props["created"] = created.isoformat()
    except Exception:
        pass
    return props


# ---------------------------------------------------------------------------
# Tracked changes — use the accepted version
# ---------------------------------------------------------------------------

def _resolve_tracked_changes(doc) -> None:
    """
    Remove revision markup so only the accepted text remains visible.

    python-docx does not expose tracked changes as a first-class API, so we
    operate directly on the underlying XML:

    - w:ins  (insertions)  — keep the content inside
    - w:del  (deletions)   — discard the content inside
    - w:rPrChange          — keep the current run properties
    """
    try:
        from lxml import etree
        body = doc.element.body

        # Remove deleted runs/paragraphs
        for node in body.findall(".//" + qn("w:del")):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)

        # Unwrap insertions (keep their children)
        for node in body.findall(".//" + qn("w:ins")):
            parent = node.getparent()
            if parent is not None:
                idx = list(parent).index(node)
                for child in list(node):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(node)

        # Strip run-property change markers
        for node in body.findall(".//" + qn("w:rPrChange")):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
    except Exception:
        # If anything fails, silently continue — the raw text is still accessible
        pass


# ---------------------------------------------------------------------------
# Core DOCX → markdown conversion
# ---------------------------------------------------------------------------

def _docx_to_sections(source: str, cfg: dict) -> tuple[str, list[dict], dict]:
    """
    Open a DOCX file and convert it to a single markdown string, a list of
    footnote/endnote definitions, and document properties.

    Returns:
        (markdown_text, footnote_definitions, doc_properties)

    footnote_definitions is a list of dicts:
        [{"id": "1", "text": "...", "kind": "footnote"}, ...]
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document(source)

    # Resolve tracked changes before processing
    if not cfg.get("preserve_tracked_changes", False):
        _resolve_tracked_changes(doc)

    # Gather footnotes, endnotes, and comments
    footnotes = _extract_footnotes(doc)
    endnotes = _extract_endnotes(doc)
    comments_map = _extract_comments(doc) if cfg.get("preserve_comments", False) else {}

    # Collect footnote/endnote text for appending at end of section
    footnote_defs: list[dict] = []
    for fn_id, fn_text in sorted(footnotes.items(), key=lambda x: int(x[0])):
        footnote_defs.append({"id": fn_id, "text": fn_text, "kind": "footnote"})
    for en_id, en_text in sorted(endnotes.items(), key=lambda x: int(x[0])):
        footnote_defs.append({"id": en_id, "text": en_text, "kind": "endnote"})

    # Build footnote id → sequential label mapping
    fn_label_map: dict[str, int] = {}
    _fn_counter = 0

    doc_props = _get_doc_properties(doc)

    # Extract images if requested
    image_paths: list[str] = []
    if cfg.get("extract_images", True):
        output_dir = cfg.get("_output_dir", "")
        if output_dir:
            assets_dir = os.path.join(output_dir, "assets")
            image_paths = _extract_images(doc, assets_dir)

    # Walk the document body in element order, handling paragraphs and tables
    md_lines: list[str] = []
    # Track footnote references encountered in runs
    fn_ref_counter = 0
    fn_ref_map: dict[str, int] = {}

    # Iterate over body children to preserve paragraph/table order
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Reconstruct paragraph from element
            para = _find_paragraph_for_element(doc, child)
            if para is not None:
                md_line, fn_ref_counter, fn_ref_map = _para_to_md_with_refs(
                    para, cfg, fn_ref_counter, fn_ref_map, comments_map
                )
                md_lines.append(md_line)
            else:
                # Fallback: extract plain text
                texts = []
                for r in child.findall(".//" + qn("w:r")):
                    for t in r.findall(qn("w:t")):
                        texts.append(t.text or "")
                md_lines.append("".join(texts))

        elif tag == "tbl":
            # Find corresponding Table object
            tbl = _find_table_for_element(doc, child)
            if tbl is not None:
                md_lines.append("")
                md_lines.append(_table_to_markdown(tbl))
                md_lines.append("")

    # Append footnote/endnote definitions using markdown footnote syntax
    if footnote_defs:
        md_lines.append("")
        for defn in footnote_defs:
            label = fn_ref_map.get(defn["id"], defn["id"])
            md_lines.append(f"[^{label}]: {defn['text']}")

    markdown_text = "\n".join(md_lines)
    return markdown_text, footnote_defs, doc_props


def _find_paragraph_for_element(doc, elem) -> object | None:
    """Find the Paragraph object corresponding to a given lxml element."""
    try:
        for para in doc.paragraphs:
            if para._element is elem:
                return para
    except Exception:
        pass
    return None


def _find_table_for_element(doc, elem) -> object | None:
    """Find the Table object corresponding to a given lxml element."""
    try:
        for tbl in doc.tables:
            if tbl._element is elem:
                return tbl
    except Exception:
        pass
    return None


def _para_to_md_with_refs(
    para,
    cfg: dict,
    fn_counter: int,
    fn_ref_map: dict,
    comments_map: dict,
) -> tuple[str, int, dict]:
    """
    Convert a paragraph to markdown, inserting footnote references inline
    and optionally embedding comments as HTML comments.

    Returns (markdown_string, updated_fn_counter, updated_fn_ref_map).
    """
    style_name = para.style.name if para.style else ""
    heading_level = _heading_level_from_style(style_name, para, cfg)

    parts: list[str] = []
    for run in para.runs:
        run_text = _run_to_markdown(run)

        # Detect footnote references in the run's XML
        for fn_ref in run._r.findall(".//" + qn("w:footnoteReference")):
            fn_id = fn_ref.get(qn("w:id"), "")
            if fn_id and fn_id not in fn_ref_map:
                fn_counter += 1
                fn_ref_map[fn_id] = fn_counter
            label = fn_ref_map.get(fn_id, fn_id)
            run_text += f"[^{label}]"

        # Detect endnote references
        for en_ref in run._r.findall(".//" + qn("w:endnoteReference")):
            en_id = en_ref.get(qn("w:id"), "")
            if en_id and en_id not in fn_ref_map:
                fn_counter += 1
                fn_ref_map[en_id] = fn_counter
            label = fn_ref_map.get(en_id, en_id)
            run_text += f"[^{label}]"

        # Detect comment references
        if comments_map:
            for cm_ref in run._r.findall(".//" + qn("w:commentReference")):
                cm_id = cm_ref.get(qn("w:id"), "")
                comment_text = comments_map.get(cm_id, "")
                if comment_text:
                    run_text += f"<!-- comment: {comment_text} -->"

        parts.append(run_text)

    text = "".join(parts)

    if heading_level is not None:
        prefix = "#" * heading_level
        return f"{prefix} {text.strip()}", fn_counter, fn_ref_map

    if "List" in style_name:
        return f"- {text}", fn_counter, fn_ref_map

    return text, fn_counter, fn_ref_map


# ---------------------------------------------------------------------------
# Chapter extraction from markdown text
# ---------------------------------------------------------------------------

def _extract_chapter_content(text: str, chapters: list) -> list:
    """
    Given detected chapter metadata (with start/end character offsets) and the
    full markdown text, attach 'content' and 'word_count' to each chapter dict.
    """
    enriched = []
    for ch in chapters:
        start = ch.get("start", 0)
        end = ch.get("end", len(text))
        raw = text[start:end]

        lines = raw.splitlines()
        title_line = ch.get("title", "").strip()

        # Strip the heading line from the body
        first_nonempty_idx = None
        for idx, line in enumerate(lines):
            if line.strip():
                first_nonempty_idx = idx
                break

        if first_nonempty_idx is not None:
            first_line = lines[first_nonempty_idx].strip()
            # Remove markdown heading markers for comparison
            first_line_text = re.sub(r"^#+\s*", "", first_line)
            if first_line_text.startswith(title_line) or first_line.startswith(title_line):
                title_line = first_line_text if first_line_text else title_line
            body_lines = lines[first_nonempty_idx + 1:]
        else:
            body_lines = lines

        content = "\n".join(body_lines).strip()

        enriched.append({
            **ch,
            "title": title_line,
            "content": content,
            "word_count": word_count(content),
        })
    return enriched


def _filter_chapters(chapters: list, skip_sections: list) -> tuple[list, list]:
    """
    Split detected chapters into special sections and regular chapters.
    Applies skip_sections filter.
    """
    special_types = {"prologue", "epilogue", "section"}
    specials = []
    regulars = []

    skip_lower = [s.lower() for s in (skip_sections or [])]

    for ch in chapters:
        title_lower = ch.get("title", "").lower()

        if any(skip in title_lower for skip in skip_lower):
            continue

        ch_type = ch.get("type", "chapter")

        if ch_type in special_types or ch.get("title", "") == "(preamble)":
            if ch.get("word_count", 0) > 0:
                specials.append(ch)
        else:
            if ch.get("word_count", 0) >= 1:
                regulars.append(ch)

    return specials, regulars


# ---------------------------------------------------------------------------
# Main conversion pipeline
# ---------------------------------------------------------------------------

def convert(source: str, output_dir: str, config: dict) -> dict:
    """
    Full conversion pipeline:
      1. Open DOCX and extract markdown text
      2. Extract images into assets/ folder
      3. Detect chapter boundaries (chapter_detector)
      4. Write output files (codex_writer)

    Returns the final result dict.
    """
    cfg = {**CONFIG, **config}
    # Store output_dir in cfg for image extraction
    cfg["_output_dir"] = output_dir

    emit({"event": "start", "source": source, "output_dir": output_dir})

    # --- Validate source ---
    if not os.path.isfile(source):
        raise FileNotFoundError(f"Source file not found: {source}")

    if not source.lower().endswith(".docx"):
        emit({"event": "warning", "message": f"Source does not have .docx extension: {source}"})

    # --- Step 1: Open and parse DOCX ---
    emit({"event": "progress", "step": 1, "message": "Opening DOCX file"})

    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed. Install it with: pip install python-docx"
        )

    markdown_text, footnote_defs, doc_props = _docx_to_sections(source, cfg)

    emit({
        "event": "progress",
        "step": 1,
        "message": (
            f"Parsed DOCX: {len(markdown_text)} characters, "
            f"{len(footnote_defs)} footnotes/endnotes"
        ),
    })

    # Derive project title: prefer document property, fall back to filename stem
    project_title = (
        doc_props.get("title", "").strip()
        or os.path.splitext(os.path.basename(source))[0]
    )
    project_author = doc_props.get("author", "")

    # --- Step 2: Detect chapter boundaries ---
    emit({"event": "progress", "step": 2, "message": "Detecting chapters"})

    hints: dict = {}
    if cfg.get("chapter_pattern"):
        hints["pattern"] = cfg["chapter_pattern"]

    detector_input = {"text": markdown_text, "hints": hints}
    detector_result = run_common("chapter_detector.py", detector_input)

    raw_chapters = detector_result.get("chapters", [])
    detection_method = detector_result.get("detection_method", "unknown")

    emit({
        "event": "progress",
        "step": 2,
        "message": f"Detected {len(raw_chapters)} sections via '{detection_method}'",
    })

    # Enrich with actual content
    enriched = _extract_chapter_content(markdown_text, raw_chapters)
    special_sections, chapters = _filter_chapters(
        enriched, skip_sections=cfg.get("skip_sections", [])
    )

    # --- Step 3: Write output ---
    emit({"event": "progress", "step": 3, "message": "Writing output files"})

    total_text_words = word_count(markdown_text)

    writer_input = {
        "output_dir": output_dir,
        "format": cfg.get("output_format", "markdown"),
        "structure": cfg.get("structure", "flat"),
        "chapters": chapters,
        "special_sections": special_sections,
        "metadata": {
            "title": project_title,
            "author": project_author,
            "source_file": source,
            "detection_method": detection_method,
            "word_count": total_text_words,
            "created": doc_props.get("created", ""),
        },
    }
    writer_result = run_common("codex_writer.py", writer_input)

    files_created = writer_result.get("files_created", 0)
    total_words = writer_result.get("total_words", 0)

    emit({"event": "progress", "step": 3, "message": f"Wrote {files_created} files"})

    result = {
        "event": "complete",
        "files": files_created,
        "total_words": total_words,
        "output_dir": output_dir,
    }
    emit(result)
    return result


# ---------------------------------------------------------------------------
# Entry point — supports both stdin JSON mode and CLI mode
# ---------------------------------------------------------------------------

def main() -> int:
    # --- CLI mode: python3 docx_converter.py source.docx output/ ---
    if len(sys.argv) >= 3:
        source = sys.argv[1]
        output_dir = sys.argv[2]
        config: dict = {}
        if len(sys.argv) >= 4:
            try:
                config = json.loads(sys.argv[3])
            except json.JSONDecodeError:
                print(
                    json.dumps({"event": "error", "message": "Third argument is not valid JSON config"}),
                    flush=True,
                )
                return 1
        try:
            convert(source, output_dir, config)
            return 0
        except Exception as exc:
            print(json.dumps({"event": "error", "message": str(exc)}), flush=True)
            return 1

    # --- Stdin JSON mode ---
    raw = sys.stdin.read()
    if not raw.strip():
        print(
            json.dumps({"event": "error", "message": "No input received on stdin"}),
            flush=True,
        )
        return 1

    try:
        payload = json.loads(raw)
    except json.JSONDecodeError as exc:
        print(json.dumps({"event": "error", "message": f"Invalid JSON input: {exc}"}), flush=True)
        return 1

    source = payload.get("source", "")
    output_dir = payload.get("output_dir", "")
    config = payload.get("config", {})

    if not source:
        print(json.dumps({"event": "error", "message": "Missing 'source' field"}), flush=True)
        return 1
    if not output_dir:
        print(json.dumps({"event": "error", "message": "Missing 'output_dir' field"}), flush=True)
        return 1

    try:
        convert(source, output_dir, config)
        return 0
    except Exception as exc:
        print(json.dumps({"event": "error", "message": str(exc)}), flush=True)
        return 1


if __name__ == "__main__":
    sys.exit(main())
